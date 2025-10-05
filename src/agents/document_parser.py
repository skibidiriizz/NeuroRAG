"""
Document Parser Agent for RAG System

This agent handles parsing various document formats including PDF, DOCX, TXT, and HTML.
It extracts text content and metadata from documents.
"""

import os
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Union
from datetime import datetime

# Document processing imports
import pypdf
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import requests
from urllib.parse import urlparse

# Local imports
from ..models.document import Document, DocumentMetadata, DocumentType, ProcessingStatus
from ..utils.logger import get_logger

logger = get_logger(__name__)


class DocumentParserAgent:
    """
    Agent responsible for parsing various document formats and extracting content.
    """
    
    def __init__(self, config: Dict[str, Any] = None):
        """
        Initialize the Document Parser Agent.
        
        Args:
            config: Configuration dictionary
        """
        self.config = config or {}
        self.max_file_size_mb = self.config.get('max_file_size_mb', 50)
        self.supported_formats = self.config.get('supported_formats', ['pdf', 'docx', 'txt', 'html'])
        self.extract_metadata = self.config.get('extract_metadata', True)
        
        logger.info("Document Parser Agent initialized")
    
    def parse_document(self, file_path: str) -> Document:
        """
        Parse a document from file path.
        
        Args:
            file_path: Path to the document file
            
        Returns:
            Document object with extracted content and metadata
            
        Raises:
            FileNotFoundError: If file doesn't exist
            ValueError: If file format is not supported
        """
        file_path = Path(file_path)
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Check file size
        file_size = file_path.stat().st_size
        if file_size > self.max_file_size_mb * 1024 * 1024:
            raise ValueError(f"File size ({file_size / (1024*1024):.1f}MB) exceeds maximum allowed size ({self.max_file_size_mb}MB)")
        
        # Determine document type
        doc_type = self._determine_document_type(file_path)
        
        if doc_type == DocumentType.UNKNOWN:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
        
        logger.info(f"Parsing {doc_type.value} document: {file_path}")
        
        try:
            # Extract content based on document type
            content = self._extract_content(file_path, doc_type)
            
            # Extract metadata if enabled
            metadata = DocumentMetadata()
            if self.extract_metadata:
                metadata = self._extract_metadata(file_path, doc_type, content)
            
            # Create document object
            document = Document(
                filename=file_path.name,
                file_path=str(file_path),
                content=content,
                document_type=doc_type,
                metadata=metadata,
                processing_status=ProcessingStatus.COMPLETED
            )
            
            logger.info(f"Successfully parsed document: {file_path.name} ({len(content)} characters)")
            return document
            
        except Exception as e:
            logger.error(f"Error parsing document {file_path}: {str(e)}")
            # Create document with error status
            document = Document(
                filename=file_path.name,
                file_path=str(file_path),
                content="",
                document_type=doc_type,
                processing_status=ProcessingStatus.FAILED,
                error_message=str(e)
            )
            return document
    
    def parse_url(self, url: str) -> Document:
        """
        Parse a document from URL (HTML only).
        
        Args:
            url: URL to fetch and parse
            
        Returns:
            Document object with extracted content
        """
        logger.info(f"Fetching document from URL: {url}")
        
        try:
            response = requests.get(url, timeout=30)
            response.raise_for_status()
            
            # Create temporary filename from URL
            parsed_url = urlparse(url)
            filename = parsed_url.path.split('/')[-1] or 'webpage.html'
            if not filename.endswith('.html'):
                filename += '.html'
            
            # Parse HTML content
            content = self._parse_html_content(response.text)
            
            # Extract metadata
            metadata = self._extract_html_metadata(response.text, url)
            
            document = Document(
                filename=filename,
                file_path=url,
                content=content,
                document_type=DocumentType.HTML,
                metadata=metadata,
                processing_status=ProcessingStatus.COMPLETED
            )
            
            logger.info(f"Successfully parsed URL: {url} ({len(content)} characters)")
            return document
            
        except Exception as e:
            logger.error(f"Error parsing URL {url}: {str(e)}")
            document = Document(
                filename=urlparse(url).path.split('/')[-1] or 'webpage.html',
                file_path=url,
                content="",
                document_type=DocumentType.HTML,
                processing_status=ProcessingStatus.FAILED,
                error_message=str(e)
            )
            return document
    
    def parse_batch(self, file_paths: List[str]) -> List[Document]:
        """
        Parse multiple documents in batch.
        
        Args:
            file_paths: List of file paths to parse
            
        Returns:
            List of Document objects
        """
        logger.info(f"Starting batch parsing of {len(file_paths)} documents")
        
        documents = []
        for file_path in file_paths:
            try:
                document = self.parse_document(file_path)
                documents.append(document)
            except Exception as e:
                logger.error(f"Failed to parse {file_path}: {str(e)}")
                # Create failed document entry
                failed_doc = Document(
                    filename=Path(file_path).name,
                    file_path=file_path,
                    content="",
                    document_type=DocumentType.UNKNOWN,
                    processing_status=ProcessingStatus.FAILED,
                    error_message=str(e)
                )
                documents.append(failed_doc)
        
        successful = len([d for d in documents if d.processing_status == ProcessingStatus.COMPLETED])
        logger.info(f"Batch parsing completed: {successful}/{len(file_paths)} documents successful")
        
        return documents
    
    def _determine_document_type(self, file_path: Path) -> DocumentType:
        """Determine document type from file extension."""
        suffix = file_path.suffix.lower()
        
        type_mapping = {
            '.pdf': DocumentType.PDF,
            '.docx': DocumentType.DOCX,
            '.doc': DocumentType.DOCX,  # Treat .doc as DOCX (limited support)
            '.txt': DocumentType.TXT,
            '.html': DocumentType.HTML,
            '.htm': DocumentType.HTML
        }
        
        return type_mapping.get(suffix, DocumentType.UNKNOWN)
    
    def _extract_content(self, file_path: Path, doc_type: DocumentType) -> str:
        """Extract text content based on document type."""
        if doc_type == DocumentType.PDF:
            return self._parse_pdf(file_path)
        elif doc_type == DocumentType.DOCX:
            return self._parse_docx(file_path)
        elif doc_type == DocumentType.TXT:
            return self._parse_txt(file_path)
        elif doc_type == DocumentType.HTML:
            return self._parse_html(file_path)
        else:
            raise ValueError(f"Unsupported document type: {doc_type}")
    
    def _parse_pdf(self, file_path: Path) -> str:
        """Parse PDF document using pypdf."""
        try:
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                text_parts = []
                
                for page_num, page in enumerate(reader.pages):
                    try:
                        text = page.extract_text()
                        if text.strip():  # Only add non-empty pages
                            text_parts.append(text)
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {page_num + 1} in {file_path}: {str(e)}")
                        continue
                
                content = "\\n\\n".join(text_parts)
                return self._clean_text(content)
                
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    def _parse_docx(self, file_path: Path) -> str:
        """Parse DOCX document using python-docx."""
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            content = "\\n\\n".join(text_parts)
            return self._clean_text(content)
            
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def _parse_txt(self, file_path: Path) -> str:
        """Parse plain text document."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()
                        return self._clean_text(content)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, try with error handling
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                content = file.read()
                return self._clean_text(content)
                
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
    
    def _parse_html(self, file_path: Path) -> str:
        """Parse HTML document using BeautifulSoup."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as file:
                html_content = file.read()
            
            return self._parse_html_content(html_content)
            
        except Exception as e:
            raise Exception(f"Error parsing HTML: {str(e)}")
    
    def _parse_html_content(self, html_content: str) -> str:
        """Parse HTML content and extract text."""
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Extract text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = "\\n".join(chunk for chunk in chunks if chunk)
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Clean extracted text."""
        if not text:
            return ""
        
        # Normalize whitespace
        text = re.sub(r'\\s+', ' ', text)
        
        # Remove excessive line breaks
        text = re.sub(r'\\n\\s*\\n\\s*\\n', '\\n\\n', text)
        
        # Remove control characters except newlines and tabs
        text = re.sub(r'[\\x00-\\x08\\x0B\\x0C\\x0E-\\x1F\\x7F]', '', text)
        
        return text.strip()
    
    def _extract_metadata(self, file_path: Path, doc_type: DocumentType, content: str) -> DocumentMetadata:
        """Extract metadata from document."""
        metadata = DocumentMetadata()
        
        # Basic file information
        stat = file_path.stat()
        metadata.file_size = stat.st_size
        metadata.creation_date = datetime.fromtimestamp(stat.st_ctime)
        metadata.modification_date = datetime.fromtimestamp(stat.st_mtime)
        
        # Content-based metadata
        metadata.word_count = len(content.split()) if content else 0
        
        # Format-specific metadata
        if doc_type == DocumentType.PDF:
            metadata = self._extract_pdf_metadata(file_path, metadata)
        elif doc_type == DocumentType.DOCX:
            metadata = self._extract_docx_metadata(file_path, metadata)
        elif doc_type == DocumentType.HTML:
            metadata = self._extract_html_metadata(content, str(file_path))
        
        # Extract potential title from filename or content
        if not metadata.title:
            # Try to get title from filename
            title_from_filename = file_path.stem.replace('_', ' ').replace('-', ' ').title()
            metadata.title = title_from_filename
        
        return metadata
    
    def _extract_pdf_metadata(self, file_path: Path, metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract PDF-specific metadata."""
        try:
            with open(file_path, 'rb') as file:
                reader = pypdf.PdfReader(file)
                
                # Page count
                metadata.page_count = len(reader.pages)
                
                # Document info
                if reader.metadata:
                    if reader.metadata.title:
                        metadata.title = reader.metadata.title
                    if reader.metadata.author:
                        metadata.author = reader.metadata.author
                    if reader.metadata.subject:
                        metadata.subject = reader.metadata.subject
                    if reader.metadata.creator:
                        metadata.extra['creator'] = reader.metadata.creator
                    if reader.metadata.producer:
                        metadata.extra['producer'] = reader.metadata.producer
                    if reader.metadata.creation_date:
                        metadata.creation_date = reader.metadata.creation_date
                    if reader.metadata.modification_date:
                        metadata.modification_date = reader.metadata.modification_date
        
        except Exception as e:
            logger.warning(f"Could not extract PDF metadata from {file_path}: {str(e)}")
        
        return metadata
    
    def _extract_docx_metadata(self, file_path: Path, metadata: DocumentMetadata) -> DocumentMetadata:
        """Extract DOCX-specific metadata."""
        try:
            doc = DocxDocument(file_path)
            core_props = doc.core_properties
            
            if core_props.title:
                metadata.title = core_props.title
            if core_props.author:
                metadata.author = core_props.author
            if core_props.subject:
                metadata.subject = core_props.subject
            if core_props.keywords:
                metadata.keywords = [kw.strip() for kw in core_props.keywords.split(',')]
            if core_props.created:
                metadata.creation_date = core_props.created
            if core_props.modified:
                metadata.modification_date = core_props.modified
            
            # Additional properties
            if core_props.category:
                metadata.extra['category'] = core_props.category
            if core_props.comments:
                metadata.extra['comments'] = core_props.comments
        
        except Exception as e:
            logger.warning(f"Could not extract DOCX metadata from {file_path}: {str(e)}")
        
        return metadata
    
    def _extract_html_metadata(self, html_content: str, source_url: str = None) -> DocumentMetadata:
        """Extract HTML-specific metadata."""
        metadata = DocumentMetadata()
        
        try:
            soup = BeautifulSoup(html_content, 'html.parser')
            
            # Title
            title_tag = soup.find('title')
            if title_tag:
                metadata.title = title_tag.get_text().strip()
            
            # Meta tags
            meta_tags = soup.find_all('meta')
            for tag in meta_tags:
                name = tag.get('name', '').lower()
                content = tag.get('content', '')
                
                if name == 'author':
                    metadata.author = content
                elif name == 'description':
                    metadata.extra['description'] = content
                elif name == 'keywords':
                    metadata.keywords = [kw.strip() for kw in content.split(',')]
                elif name == 'language' or name == 'lang':
                    metadata.language = content
            
            # Open Graph tags
            og_tags = soup.find_all('meta', property=lambda x: x and x.startswith('og:'))
            for tag in og_tags:
                prop = tag.get('property', '')
                content = tag.get('content', '')
                
                if prop == 'og:title' and not metadata.title:
                    metadata.title = content
                elif prop == 'og:description':
                    metadata.extra['og_description'] = content
            
            if source_url:
                metadata.source_url = source_url
        
        except Exception as e:
            logger.warning(f"Could not extract HTML metadata: {str(e)}")
        
        return metadata
    
    def get_supported_formats(self) -> List[str]:
        """Get list of supported document formats."""
        return self.supported_formats.copy()
    
    def validate_file(self, file_path: str) -> Dict[str, Any]:
        """
        Validate if a file can be processed.
        
        Args:
            file_path: Path to file
            
        Returns:
            Dictionary with validation results
        """
        file_path = Path(file_path)
        result = {
            'valid': False,
            'errors': [],
            'warnings': [],
            'file_info': {}
        }
        
        # Check if file exists
        if not file_path.exists():
            result['errors'].append(f"File does not exist: {file_path}")
            return result
        
        # Check file size
        file_size = file_path.stat().st_size
        result['file_info']['size_bytes'] = file_size
        result['file_info']['size_mb'] = file_size / (1024 * 1024)
        
        if file_size > self.max_file_size_mb * 1024 * 1024:
            result['errors'].append(f"File size ({result['file_info']['size_mb']:.1f}MB) exceeds maximum allowed size ({self.max_file_size_mb}MB)")
        
        # Check format support
        doc_type = self._determine_document_type(file_path)
        result['file_info']['detected_type'] = doc_type.value
        
        if doc_type == DocumentType.UNKNOWN:
            result['errors'].append(f"Unsupported file format: {file_path.suffix}")
        elif doc_type.value not in self.supported_formats:
            result['errors'].append(f"File format not enabled: {doc_type.value}")
        
        # Check file readability
        try:
            with open(file_path, 'rb') as f:
                f.read(1)  # Try to read first byte
        except PermissionError:
            result['errors'].append("Permission denied reading file")
        except Exception as e:
            result['errors'].append(f"Cannot read file: {str(e)}")
        
        result['valid'] = len(result['errors']) == 0
        return result